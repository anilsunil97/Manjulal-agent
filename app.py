import os
import io
import base64
import time
import hashlib
import json
import pandas as pd
import streamlit as st
from docx import Document as DocxDocument
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain_huggingface import HuggingFaceEmbeddings

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Load environment variables
load_dotenv()

def get_secret(key: str) -> str:
    value = os.getenv(key)
    if not value:
        try:
            value = st.secrets[key]
        except (KeyError, FileNotFoundError):
            pass
    return value or ""

groq_api_key = get_secret("GROQ_API_KEY")

DOCS_DIR    = os.path.join(os.path.dirname(__file__), "petpooja_docs")
INDEX_DIR   = os.path.join(os.path.dirname(__file__), "faiss_index")
HASH_FILE   = os.path.join(INDEX_DIR, "docs_hash.json")

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(page_title="Manjulal Agent", page_icon="💃", layout="centered")

# ── Header ────────────────────────────────────────────────────────────────────
def load_as_base64(path: str) -> str:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()

col_logo, col_title = st.columns([1, 5])
with col_logo:
    gif_path = os.path.join(os.path.dirname(__file__), "static", "motu.gif")
    mp4_path = os.path.join(os.path.dirname(__file__), "static", "WhatsApp Video 2026-07-28 at 20.12.36.mp4")
    if os.path.exists(gif_path):
        b64 = load_as_base64(gif_path)
        st.markdown(f'<img src="data:image/gif;base64,{b64}" width="80" style="border-radius:12px;">',
                    unsafe_allow_html=True)
    elif os.path.exists(mp4_path):
        b64 = load_as_base64(mp4_path)
        st.markdown(f'<video width="80" autoplay loop muted playsinline style="border-radius:12px;">'
                    f'<source src="data:video/mp4;base64,{b64}" type="video/mp4"></video>',
                    unsafe_allow_html=True)
    else:
        st.markdown('<img src="https://media.giphy.com/media/l3vR4yk0X20KimqJ2/giphy.gif" width="80" style="border-radius:12px;">',
                    unsafe_allow_html=True)
with col_title:
    st.title("Manjulal Agent")

# ── LLM ───────────────────────────────────────────────────────────────────────
llm = ChatGroq(groq_api_key=groq_api_key, model_name="openai/gpt-oss-120b")

# ── Prompts ───────────────────────────────────────────────────────────────────
doc_prompt = ChatPromptTemplate.from_template(
    """You are a Petpooja support assistant. Answer the question precisely and concisely using ONLY the context below.

Rules:
- Give a direct, specific answer. No filler or preamble.
- Use bullet points or numbered steps only when the answer is a process or list.
- If multiple relevant details exist, include them briefly.
- If the answer is not in the context, respond with exactly: "We are updating more information."
- Do NOT use general knowledge or make up answers.

Context:
{context}

Question: {input}

Answer:"""
)

chat_prompt = ChatPromptTemplate.from_messages([
    ("system", """You are a Petpooja support assistant. Your sole purpose is to help users with Petpooja dashboard functionality and pricing.

Rules:
1. For greetings (hello, hi, good morning, how are you, etc.) reply briefly and friendly, then suggest the user ask about Petpooja dashboard features or pricing. Example: "Hi there! Feel free to ask me about Petpooja dashboard functionality or pricing."
2. For ANY other topic that is not related to Petpooja dashboard functionality or pricing, respond with exactly:
   "I'm here to help with Petpooja dashboard functionality and pricing. Please ask me something related to that!"
3. Do NOT answer general knowledge, coding, weather, or any off-topic questions.
4. Keep all responses concise and on-topic."""),
    ("human", "{input}"),
])

def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)


# ── File loaders ──────────────────────────────────────────────────────────────
def load_pdf(filepath):
    """Load PDF using PyPDFLoader; fall back to pdfplumber for scanned/complex PDFs."""
    docs = PyPDFLoader(filepath).load()
    # If all pages came back empty, try pdfplumber
    if PDFPLUMBER_AVAILABLE and all(not d.page_content.strip() for d in docs):
        plumber_docs = []
        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                # Also extract any tables on the page
                for table in page.extract_tables():
                    rows = []
                    for row in table:
                        rows.append(" | ".join(cell or "" for cell in row))
                    text += "\n" + "\n".join(rows)
                if text.strip():
                    plumber_docs.append(Document(
                        page_content=text.strip(),
                        metadata={"source": os.path.basename(filepath), "page": i}
                    ))
        return plumber_docs if plumber_docs else docs
    return docs

def load_excel(filepath):
    docs = []
    xl = pd.ExcelFile(filepath)
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        text = f"Sheet: {sheet}\n\n{df.to_string(index=False)}"
        docs.append(Document(page_content=text,
                             metadata={"source": os.path.basename(filepath), "sheet": sheet}))
    return docs

def load_docx(filepath):
    """Load docx extracting both paragraphs and table content."""
    doc = DocxDocument(filepath)
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Extract tables — each row becomes a pipe-separated line
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    return [Document(page_content=text, metadata={"source": os.path.basename(filepath)})]

def load_all_docs_from_folder(folder):
    all_docs = []
    supported = (".pdf", ".xlsx", ".xls", ".docx")
    for filename in os.listdir(folder):
        # Skip Office temp/lock files (e.g. ~$filename.docx)
        if filename.startswith("~$"):
            continue
        if not filename.lower().endswith(supported):
            continue
        filepath = os.path.join(folder, filename)
        ext = os.path.splitext(filename)[1].lower()
        try:
            if ext == ".pdf":
                all_docs.extend(load_pdf(filepath))
            elif ext in (".xlsx", ".xls"):
                all_docs.extend(load_excel(filepath))
            elif ext == ".docx":
                all_docs.extend(load_docx(filepath))
        except Exception as e:
            st.warning(f"⚠️ Could not load `{filename}`: {e}")
    return all_docs


# ── Hash helpers — detect if docs folder changed ──────────────────────────────
def compute_docs_hash(folder):
    """MD5 hash of all filenames + modification times in the folder."""
    supported = (".pdf", ".xlsx", ".xls", ".docx")
    entries = sorted([
        (f, os.path.getmtime(os.path.join(folder, f)))
        for f in os.listdir(folder)
        if f.lower().endswith(supported)
    ])
    return hashlib.md5(json.dumps(entries).encode()).hexdigest()

def load_saved_hash():
    if os.path.exists(HASH_FILE):
        with open(HASH_FILE) as f:
            return json.load(f).get("hash", "")
    return ""

def save_hash(h):
    os.makedirs(INDEX_DIR, exist_ok=True)
    with open(HASH_FILE, "w") as f:
        json.dump({"hash": h}, f)


# ── Embeddings model (cached so it loads only once) ───────────────────────────
@st.cache_resource(show_spinner="Loading embedding model…")
def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name="all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


# ── Load or build vector store ────────────────────────────────────────────────
def build_and_save_index(docs_dir):
    embeddings = get_embeddings()
    all_docs = load_all_docs_from_folder(docs_dir)
    if not all_docs:
        return None, 0

    # Smaller chunks = more focused retrieval = more precise answers
    splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=100)
    split_docs = splitter.split_documents(all_docs)
    total = len(split_docs)

    progress = st.progress(0, text="Embedding chunks…")
    BATCH_SIZE = 100
    vectorstore = None
    for i in range(0, total, BATCH_SIZE):
        batch = split_docs[i: i + BATCH_SIZE]
        if vectorstore is None:
            vectorstore = FAISS.from_documents(batch, embeddings)
        else:
            vectorstore.add_documents(batch)
        progress.progress(min((i + BATCH_SIZE) / total, 1.0),
                          text=f"Embedding {min(i + BATCH_SIZE, total)}/{total} chunks…")
    progress.empty()

    # Save index to disk — will be reused on next startup if docs haven't changed
    os.makedirs(INDEX_DIR, exist_ok=True)
    vectorstore.save_local(INDEX_DIR)
    save_hash(compute_docs_hash(docs_dir))
    return vectorstore, total

def load_index_from_disk():
    embeddings = get_embeddings()
    return FAISS.load_local(INDEX_DIR, embeddings, allow_dangerous_deserialization=True)


# ── Auto-load index on startup if it exists and docs haven't changed ──────────
if not os.path.isdir(DOCS_DIR):
    os.makedirs(DOCS_DIR)

files_present = [f for f in os.listdir(DOCS_DIR)
                 if f.lower().endswith((".pdf", ".xlsx", ".xls", ".docx"))
                 and not f.startswith("~$")]

index_exists = os.path.exists(os.path.join(INDEX_DIR, "index.faiss"))

if "vectors" not in st.session_state and index_exists:
    current_hash = compute_docs_hash(DOCS_DIR) if files_present else ""
    saved_hash   = load_saved_hash()
    if current_hash == saved_hash:
        # Docs unchanged — load saved index silently, no re-chunking needed
        try:
            st.session_state.vectors   = load_index_from_disk()
            st.session_state.doc_count = "cached"
        except Exception:
            pass

# ── Session state ─────────────────────────────────────────────────────────────
if "messages" not in st.session_state:
    st.session_state.messages = []

# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("📁 Document Mode")

    # ── Upload new documents ──────────────────────────────────────────────────
    st.subheader("📤 Upload Documents")
    uploaded_files = st.file_uploader(
        "Add PDF, DOCX, or Excel files",
        type=["pdf", "docx", "xlsx", "xls"],
        accept_multiple_files=True,
        help="Files are saved to petpooja_docs and indexed automatically."
    )
    if uploaded_files:
        saved_names = []
        for uf in uploaded_files:
            save_path = os.path.join(DOCS_DIR, uf.name)
            with open(save_path, "wb") as f:
                f.write(uf.getbuffer())
            saved_names.append(uf.name)
        st.success(f"✅ Saved: {', '.join(saved_names)}\nClick Re-Embed to index them.")
        st.rerun()

    st.divider()

    docs_changed = files_present and (compute_docs_hash(DOCS_DIR) != load_saved_hash())

    # Only show embed button if docs changed or no index exists yet
    if docs_changed or not index_exists:
        embed_btn = st.button(
            "🔄 Re-Embed Documents" if (docs_changed and index_exists) else "📥 Build Index",
            use_container_width=True,
            help="Docs have changed — re-embed recommended" if docs_changed else "Build the search index from your documents"
        )
    else:
        embed_btn = False
        st.info("✅ Index is up to date. No re-embedding needed.")

    if "vectors" in st.session_state:
        label = st.session_state.get("doc_count", "")
        if label == "cached":
            st.success("✅ Loaded from saved index")
        else:
            st.success(f"✅ {label} chunks embedded")
        if docs_changed:
            st.warning("⚠️ Documents changed — click Re-Embed")
    else:
        st.info("No index loaded yet.\nChat works in general mode.")

    if st.button("🗑️ Clear Chat", use_container_width=True):
        st.session_state.messages = []
        st.rerun()

    if st.button("🔄 Reset Index", use_container_width=True):
        for key in ["vectors", "doc_count"]:
            st.session_state.pop(key, None)
        st.rerun()

# ── Embed on button click ─────────────────────────────────────────────────────
if embed_btn:
    if not files_present:
        st.error("No files found in `petpooja_docs`.")
    else:
        try:
            vectorstore, total = build_and_save_index(DOCS_DIR)
            if vectorstore:
                st.session_state.vectors   = vectorstore
                st.session_state.doc_count = total
                st.rerun()
        except Exception as e:
            st.error(f"❌ Embedding failed: {e}")

# ── Chat history ──────────────────────────────────────────────────────────────
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# ── Chat input ────────────────────────────────────────────────────────────────
if question := st.chat_input("Ask me anything…"):
    st.session_state.messages.append({"role": "user", "content": question})
    with st.chat_message("user"):
        st.markdown(question)

    with st.chat_message("assistant"):
        with st.spinner("Thinking…"):
            try:
                if "vectors" in st.session_state:
                    # MMR retrieval: fetch 6 diverse, relevant chunks for precise answers
                    retriever = st.session_state.vectors.as_retriever(
                        search_type="mmr",
                        search_kwargs={"k": 6, "fetch_k": 20, "lambda_mult": 0.7}
                    )
                    rag_chain = (
                        {"context": retriever | format_docs, "input": RunnablePassthrough()}
                        | doc_prompt | llm | StrOutputParser()
                    )
                    answer = rag_chain.invoke(question)
                else:
                    chain  = chat_prompt | llm | StrOutputParser()
                    answer = chain.invoke({"input": question})

                st.markdown(answer)
                st.session_state.messages.append({"role": "assistant", "content": answer})
            except Exception as e:
                st.error(f"❌ Error: {e}")
